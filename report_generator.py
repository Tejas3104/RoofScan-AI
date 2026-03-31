from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os


SECTION_HEADINGS = [
    "1. PROPERTY ISSUE SUMMARY",
    "2. AREA-WISE OBSERVATIONS",
    "3. PROBABLE ROOT CAUSE",
    "4. SEVERITY ASSESSMENT",
    "5. RECOMMENDED ACTIONS",
    "6. ADDITIONAL NOTES",
    "7. MISSING OR UNCLEAR INFORMATION",
]

# All RGBColor values must be integers 0-255
SEVERITY_COLORS = {
    "high":   RGBColor(192, 0,   0),    # Red
    "medium": RGBColor(255, 128, 0),    # Orange
    "low":    RGBColor(0,   112, 192),  # Blue
}

# Brand colors
COLOR_BRAND_DARK  = RGBColor(31,  56,  100)   # #1F3864
COLOR_BRAND_MID   = RGBColor(46,  116, 181)   # #2E74B5
COLOR_GREY_LIGHT  = RGBColor(112, 112, 112)   # #707070
COLOR_GREY_DATE   = RGBColor(144, 144, 144)   # #909090
COLOR_FOOTER      = RGBColor(160, 160, 160)   # #A0A0A0


def add_horizontal_line(doc):
    """Add a horizontal rule to the document."""
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def style_severity_paragraph(para, text):
    """Color severity text based on High/Medium/Low."""
    lower = text.lower()
    color = None
    if "high" in lower:
        color = SEVERITY_COLORS["high"]
    elif "medium" in lower:
        color = SEVERITY_COLORS["medium"]
    elif "low" in lower:
        color = SEVERITY_COLORS["low"]

    if color:
        for run in para.runs:
            run.font.color.rgb = color


def create_ddr_report(ddr_text: str, image_paths: list, output_path: str = "outputs/DDR_Report.docx") -> str:
    """
    Convert DDR text + images into a professional Word document.
    Returns path to saved .docx file.
    """
    os.makedirs("outputs", exist_ok=True)
    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Cover Header ──
    title_para = doc.add_heading("", level=0)
    title_run = title_para.add_run("Detailed Diagnostic Report")
    title_run.font.size = Pt(26)
    title_run.font.color.rgb = COLOR_BRAND_DARK
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("Urbanroof Pvt Ltd — AI-Generated Property Report")
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = COLOR_GREY_LIGHT
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f"Generated on: {datetime.now().strftime('%B %d, %Y  %I:%M %p')}")
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = COLOR_GREY_DATE
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_horizontal_line(doc)
    doc.add_paragraph("")

    # ── Parse and write DDR sections ──
    lines = ddr_text.split("\n")
    current_section = None
    image_index = 0

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Detect main section headings
        is_section = False
        for heading in SECTION_HEADINGS:
            if line.upper().startswith(heading[:6]):
                is_section = True
                current_section = line
                h = doc.add_heading(line, level=1)
                for run in h.runs:
                    run.font.color.rgb = COLOR_BRAND_DARK
                add_horizontal_line(doc)

                # Place images after Area-wise Observations
                if "AREA-WISE" in line.upper() and image_index < len(image_paths):
                    _insert_images(doc, image_paths, image_index, max_images=3)
                    image_index += 3
                break

        if is_section:
            continue

        # Sub-headings (short lines ending with ":")
        if line.endswith(":") and len(line) < 80:
            h2 = doc.add_heading(line, level=2)
            for run in h2.runs:
                run.font.size = Pt(12)
                run.font.color.rgb = COLOR_BRAND_MID
            continue

        # Bullet points
        if line.startswith("-") or line.startswith("•") or line.startswith("*"):
            clean = line.lstrip("-•* ").strip()
            para = doc.add_paragraph(clean, style="List Bullet")
            para.runs[0].font.size = Pt(10.5)
            if current_section and "SEVERITY" in current_section.upper():
                style_severity_paragraph(para, clean)
            continue

        # Numbered list items
        if len(line) > 2 and line[0].isdigit() and line[1] in ".)":
            para = doc.add_paragraph(line, style="List Number")
            para.runs[0].font.size = Pt(10.5)
            continue

        # Normal paragraph
        para = doc.add_paragraph(line)
        if para.runs:
            para.runs[0].font.size = Pt(10.5)
            if current_section and "SEVERITY" in current_section.upper():
                style_severity_paragraph(para, line)

    # ── Remaining images at end ──
    if image_index < len(image_paths):
        doc.add_page_break()
        img_heading = doc.add_heading("Supporting Images", level=1)
        for run in img_heading.runs:
            run.font.color.rgb = COLOR_BRAND_DARK
        add_horizontal_line(doc)
        _insert_images(doc, image_paths, image_index, max_images=len(image_paths))

    # ── Footer ──
    doc.add_paragraph("")
    add_horizontal_line(doc)
    footer_para = doc.add_paragraph()
    footer_run = footer_para.add_run(
        "This report was auto-generated by Urbanroof AI System. "
        "All findings are based solely on the provided inspection and thermal documents."
    )
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = COLOR_FOOTER
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(output_path)
    return output_path


def _insert_images(doc, image_paths, start_index, max_images=3):
    """Insert images into document with captions."""
    count = 0
    for i in range(start_index, len(image_paths)):
        if count >= max_images:
            break
        img_path = image_paths[i]
        try:
            doc.add_picture(img_path, width=Inches(5.0))
            caption = doc.add_paragraph(f"Figure: {os.path.basename(img_path)}")
            caption.runs[0].font.size = Pt(9)
            caption.runs[0].font.italic = True
            caption.runs[0].font.color.rgb = COLOR_GREY_LIGHT
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph("")
            count += 1
        except Exception as e:
            doc.add_paragraph(f"[Image not available: {os.path.basename(img_path)}]")
